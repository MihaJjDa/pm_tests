from docx import Document
from docx.shared import Inches
import random
from xlrd import open_workbook
from xlwt import Workbook, Font

def add_question(txt, qst, num, images, document, sheet, ws):
    txt = txt.split('\n')
    ok = False
    i = 0
    while not ok and i < len(txt):
        if txt[i].find(str(qst) + '.') >= 0:
            ok = True
        else:
            i = i + 1
    document.add_paragraph(str(num)+txt[i][txt[i].find('.'):])
    i = i + 1
    while txt[i] != '':
        document.add_paragraph(txt[i][2:], style='List Bullet')
        i = i + 1
    if qst in images:
        document.add_picture('data\\' + str(qst) + '.png', width=Inches(5))
    document.add_paragraph('', style=None)
    row = sheet.row_values(qst-1)
    ws.write(num-1, 0, num)
    ws.write(num-1, 1, qst)
    ws.write(num-1, 2, row[0])



print('Программа перемешивания вопросов')
print('Пучкин Д.А., danila.puchkin@mail.ru')
print('Март 2019 года')
print('----------------------------------------')
print('Вопросы должны находиться в каталоге data в файле data/questions.txt')
print('----------------------------------------')
print('Формат вопросов: "<Номер вопроса>. <Вопрос>"')
print('Вопросы нумеруются с 1')
print('----------------------------------------')
print('Формат ответов: произвольное количество строк в формате "* <Ответа>"')
print('Вопросы должны отделяться пустой строкой')
print('----------------------------------------')
print('Для приложения изображения к вопросу необходимо')
print('поместить в каталог data изображение с именем "<Номер вопроса>.png"')
print('----------------------------------------')
print('Ответы должны находиться в каталоге data в файле data/answers.xlsx')
print('----------------------------------------')
print('Формат ответов: первый столбец с ответами')
print('Количество ответов должно совпадать с количеством вопросов')
print('----------------------------------------')
print('В случае некорректной подачи данных корректная работы программы не гарантируется')
print('Для начала работы нажмите Enter')
a = input()

print('----------------------------------------')
print('Введите количество вопросов')
print('Нажмите Enter для использования настроек по умолчанию (188 вопросов)')
l = input()
if l == '':
    l = list(range(1, 189))
else:
    l = int(l)
ll = len(l)

print('----------------------------------------')
print('Введите номера вопросов с изображениями в одну строку через пробел')
print('Введите 0, если изображений нет')
print('Нажмите Enter для использования настроек по умолчанию')
print('Вопросы с изображениями по умолчанию:')
print('97 103 105 107 115 127 137 149 153 155 161 164 186')
images = input()
if images == '0':
    images = []
elif images == '':
    images = [97, 103, 105, 107, 115, 127, 137, 149, 153, 155, 161, 164, 186]
else:
    images = [int(i) for i in images.split(' ')]


ok = False
while not ok:
    print('----------------------------------------')
    print('Введите параметр перемешивания - проивзольное число')
    print('Нажмите Enter, чтобы пропустить')
    seed = input()
    if seed == '':
        random.seed()
        ok = True
    elif seed.isdigit():
        seed = int(seed)
        random.seed(seed)
        ok = True
    else:
        print('Неверный ввод! Введите произвольное число или нажмите Enter, чтобы пропустить')


document = Document()

rb = open_workbook('data\\answers.xlsx')
sheet = rb.sheet_by_index(0)

font0 = Font()
font0.name = 'Times New Roman'
font0.colour_index = 2
font0.bold = True
wb = Workbook()
ws = wb.add_sheet('Лист1')

l = random.sample(l, ll)

f = open('data\\questions.txt', 'r').read()
for n, q in enumerate(l):
    add_question(f, q, n+1, images, document, sheet, ws)

document.save('questions.docx')
wb.save('new_answers.xls')

print('Файл с вопросами: questions.docx')
print('Файл с ответами:  new_answers.xls')
print('Нажмите Enter')
a = input()